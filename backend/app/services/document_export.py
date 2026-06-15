import base64
import html
import io
import mimetypes
import zipfile
from datetime import datetime, timezone

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from fpdf import FPDF

from app.schemas import BatchDocument, ImageAsset, PageDocument
from app.services.storage import image_storage

SUPPORTED_FORMATS = {"docx", "pdf"}


def build_page_document_html(page: PageDocument) -> str:
    seo = page.seo
    title = seo.title if seo else f"{page.puja} in {page.city}"
    description = seo.description if seo else ""
    keywords = ", ".join(seo.keywords) if seo and seo.keywords else ""
    focus = seo.focus_keyword if seo else ""
    tagline = seo.tagline if seo else ""
    breadcrumb = seo.breadcrumb if seo else ""

    faq_html = ""
    if page.faq:
        faq_items = "".join(
            f"<div class='faq-item'><h3>{html.escape(item.question)}</h3>"
            f"<p>{html.escape(item.answer)}</p></div>"
            for item in page.faq
        )
        faq_html = f"<section class='faq'><h2>Frequently Asked Questions</h2>{faq_items}</section>"

    images_html = ""
    if page.images:
        figures = "".join(_image_figure(img) for img in page.images)
        images_html = f"<section class='images'><h2>Images</h2>{figures}</section>"

    seo_appendix = f"""
    <section class='seo-appendix'>
      <h2>SEO Metadata</h2>
      <table>
        <tr><th>Meta Title</th><td>{html.escape(title)}</td></tr>
        <tr><th>Meta Description</th><td>{html.escape(description)}</td></tr>
        <tr><th>Focus Keyword</th><td>{html.escape(focus)}</td></tr>
        <tr><th>Keywords</th><td>{html.escape(keywords)}</td></tr>
        <tr><th>URL Slug</th><td>/{html.escape(page.slug)}/</td></tr>
        <tr><th>Tagline</th><td>{html.escape(tagline)}</td></tr>
        <tr><th>Breadcrumb</th><td>{html.escape(breadcrumb)}</td></tr>
      </table>
    </section>
    """

    generated = page.generated_at or datetime.now(timezone.utc)
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <title>{html.escape(title)}</title>
  <style>
    body {{ font-family: Georgia, serif; margin: 2rem; color: #3d2918; line-height: 1.65; }}
    h1, h2, h3 {{ color: #5c2e1a; }}
    .lead {{ font-size: 1.1rem; }}
    .meta {{ font-size: 0.85rem; color: #8a7a6a; margin-bottom: 2rem; }}
    .faq-item {{ margin: 0.75rem 0; padding: 0.75rem; border-left: 4px solid #c87828; }}
    .images img {{ max-width: 100%; margin: 1rem 0; }}
    table {{ width: 100%; border-collapse: collapse; }}
    th, td {{ border: 1px solid #ccc; padding: 0.5rem; text-align: left; }}
  </style>
</head>
<body>
  <p class="meta">Naman Puja · Generated {generated.strftime('%Y-%m-%d %H:%M UTC')}</p>
  <article>{page.content}</article>
  {images_html}
  {faq_html}
  {seo_appendix}
</body>
</html>"""


def build_page_docx(page: PageDocument) -> bytes:
    document = Document()
    title = page.seo.title if page.seo else f"{page.puja} in {page.city}"
    document.add_heading(title, 0)
    document.add_paragraph(
        f"Generated for {page.puja} · {page.city}, {page.state}, {page.country}"
    ).italic = True

    _append_html_to_docx(document, page.content)

    if page.images:
        document.add_heading("Images", level=1)
        for image in page.images:
            local_path = image_storage.resolve_local_path(image.path)
            if local_path and local_path.exists():
                document.add_picture(str(local_path), width=Inches(5.8))
            caption = document.add_paragraph(image.caption or image.alt)
            caption.runs[0].font.size = Pt(10)

    if page.faq:
        document.add_heading("Frequently Asked Questions", level=1)
        for item in page.faq:
            document.add_heading(item.question, level=2)
            document.add_paragraph(item.answer)

    if page.seo:
        document.add_heading("SEO Metadata", level=1)
        for label, value in [
            ("Meta Title", page.seo.title),
            ("Meta Description", page.seo.description),
            ("Focus Keyword", page.seo.focus_keyword),
            ("Keywords", ", ".join(page.seo.keywords)),
            ("URL Slug", f"/{page.slug}/"),
            ("Tagline", page.seo.tagline),
            ("Breadcrumb", page.seo.breadcrumb),
        ]:
            document.add_paragraph(f"{label}: {value}")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_page_pdf(page: PageDocument) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(18, 18, 18)

    title = page.seo.title if page.seo else f"{page.puja} in {page.city}"
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, _pdf_text(title))
    pdf.ln(4)
    pdf.set_font("Helvetica", size=10)
    pdf.set_text_color(120, 100, 80)
    pdf.multi_cell(0, 6, _pdf_text(f"{page.puja} · {page.city}, {page.state}, {page.country}"))
    pdf.set_text_color(0, 0, 0)
    pdf.ln(6)

    _append_html_to_pdf(pdf, page.content)

    if page.images:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, "Images", new_x="LMARGIN", new_y="NEXT")
        for image in page.images:
            local_path = image_storage.resolve_local_path(image.path)
            if local_path and local_path.exists():
                pdf.image(str(local_path), w=170)
                pdf.ln(4)
            caption = image.caption or image.alt
            if caption:
                pdf.set_font("Helvetica", "I", 10)
                pdf.multi_cell(0, 6, _pdf_text(caption))
                pdf.ln(4)

    if page.faq:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, "Frequently Asked Questions", new_x="LMARGIN", new_y="NEXT")
        for item in page.faq:
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 8, _pdf_text(item.question))
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", size=11)
            pdf.multi_cell(0, 6, _pdf_text(item.answer))
            pdf.ln(4)

    if page.seo:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, "SEO Metadata", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", size=11)
        for label, value in [
            ("Meta Title", page.seo.title),
            ("Meta Description", page.seo.description),
            ("Focus Keyword", page.seo.focus_keyword),
            ("Keywords", ", ".join(page.seo.keywords)),
            ("URL Slug", f"/{page.slug}/"),
            ("Tagline", page.seo.tagline),
            ("Breadcrumb", page.seo.breadcrumb),
        ]:
            pdf.multi_cell(0, 6, _pdf_text(f"{label}: {value}"))
            pdf.ln(2)

    return pdf.output()


def build_page_export(page: PageDocument, export_format: str) -> tuple[bytes, str, str]:
    normalized = export_format.lower()
    if normalized not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported format: {export_format}")

    if normalized == "docx":
        return build_page_docx(page), f"{page.slug}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return build_page_pdf(page), f"{page.slug}.pdf", "application/pdf"


def build_batch_zip(batch: BatchDocument, pages: list[PageDocument], export_format: str = "pdf") -> bytes:
    normalized = export_format.lower()
    if normalized not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported format: {export_format}")

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        for page in pages:
            data, filename, _ = build_page_export(page, normalized)
            archive.writestr(filename, data)
            _add_images_to_zip(archive, page)
    return buffer.getvalue()


def _add_images_to_zip(archive: zipfile.ZipFile, page: PageDocument) -> None:
    folder = f"{page.slug}/images"
    for image in page.images:
        local_path = image_storage.resolve_local_path(image.path)
        if local_path and local_path.exists():
            archive.write(local_path, arcname=f"{folder}/{local_path.name}")


def _pdf_text(value: str) -> str:
    return value.encode("latin-1", errors="replace").decode("latin-1")


def _append_html_to_pdf(pdf: FPDF, html_content: str) -> None:
    soup = BeautifulSoup(html_content, "html.parser")
    for element in soup.find_all(["h1", "h2", "h3", "p", "li"]):
        text = element.get_text(strip=True)
        if not text:
            continue
        pdf.set_x(pdf.l_margin)
        if element.name == "h1":
            pdf.set_font("Helvetica", "B", 15)
            pdf.multi_cell(0, 9, _pdf_text(text))
        elif element.name == "h2":
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 8, _pdf_text(text))
        elif element.name == "h3":
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, 7, _pdf_text(text))
        elif element.name == "li":
            pdf.set_font("Helvetica", size=11)
            pdf.multi_cell(0, 6, _pdf_text(f"- {text}"))
        else:
            pdf.set_font("Helvetica", size=11)
            pdf.multi_cell(0, 6, _pdf_text(text))
        pdf.ln(2)


def _append_html_to_docx(document: Document, html_content: str) -> None:
    soup = BeautifulSoup(html_content, "html.parser")
    for element in soup.find_all(["h1", "h2", "h3", "p", "ul", "li"]):
        text = element.get_text(strip=True)
        if not text:
            continue
        if element.name == "h1":
            document.add_heading(text, level=1)
        elif element.name == "h2":
            document.add_heading(text, level=2)
        elif element.name == "h3":
            document.add_heading(text, level=3)
        elif element.name == "li":
            document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(text)


def _image_figure(img: ImageAsset) -> str:
    src = _image_src_for_export(img.path)
    return (
        f"<figure><img src='{src}' alt='{html.escape(img.alt)}'/>"
        f"<figcaption>{html.escape(img.caption or img.alt)}</figcaption></figure>"
    )


def _image_src_for_export(public_path: str) -> str:
    local_path = image_storage.resolve_local_path(public_path)
    if not local_path or not local_path.exists():
        return html.escape(public_path)

    mime, _ = mimetypes.guess_type(local_path.name)
    if not mime or not mime.startswith("image/"):
        mime = "image/jpeg"
    encoded = base64.b64encode(local_path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{encoded}"
